#!/usr/bin/env python3
"""Build the Gemini-ready Classroom Creation Request handbook from Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "6161FF"
ACCENT_DARK = "3F3FBF"
ACCENT_LIGHT = "EEEEFF"
INK = "292F3D"
MUTED = "676879"
LINE = "DADCE0"
SOFT = "F6F7FB"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def table_widths(rows):
    cols = len(rows[0])
    if cols == 2:
        return [2700, 6660]
    if cols == 3:
        return [1900, 2300, 5160]
    if cols == 4:
        return [1300, 1900, 2400, 3760]
    return [PAGE_WIDTH_DXA // cols] * cols


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color=ACCENT, size="18", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    headings = [
        ("Heading 1", 16, 18, 10, ACCENT_DARK),
        ("Heading 2", 13, 14, 7, ACCENT_DARK),
        ("Heading 3", 12, 10, 5, ACCENT_DARK),
    ]
    for name, size, before, after, color in headings:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in [s.name for s in styles]:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05

    if "Contents Entry" not in [s.name for s in styles]:
        contents = styles.add_style("Contents Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        contents = styles["Contents Entry"]
    contents.font.name = "Calibri"
    contents.font.size = Pt(10.5)
    contents.font.color.rgb = RGBColor.from_string(INK)
    contents.paragraph_format.space_after = Pt(3)
    contents.paragraph_format.left_indent = Inches(0.2)


def add_abstract_num(numbering, abstract_id, fmt, text):
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    next_id = max(abstract_ids or [0]) + 1
    bullet_id = next_id
    decimal_id = next_id + 1
    add_abstract_num(numbering, bullet_id, "bullet", "•")
    add_abstract_num(numbering, decimal_id, "decimal", "%1.")
    return bullet_id, decimal_id


def new_num_id(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text, size=None, color=INK):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=size or 11, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size or 11, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Courier New", size=(size or 11) - 1, color=ACCENT_DARK)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=size or 11, color=color)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_single_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, table_widths(rows))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    font_size = 8.5 if cols >= 4 else 9 if cols == 3 else 9.5
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, ACCENT_LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline(p, rows[r_idx][c_idx], size=font_size)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    if len(rows) <= 16:
        add_single_table(doc, rows)
        return
    header = rows[0]
    body = rows[1:]
    for start in range(0, len(body), 10):
        if start:
            label = doc.add_paragraph()
            label.paragraph_format.space_before = Pt(4)
            label.paragraph_format.space_after = Pt(4)
            label.paragraph_format.keep_with_next = True
            run = label.add_run("Column mapping continued")
            set_font(run, size=8.5, color=MUTED, italic=True)
        add_single_table(doc, [header] + body[start:start + 10])


def shade_paragraph(paragraph, fill=SOFT, border=LINE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def build_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SYSTEM KNOWLEDGE BASE")
    set_font(r, size=10, color=ACCENT, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("Classroom Creation Request")
    set_font(r, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("System knowledge base and operations handbook")
    set_font(r, size=15, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(24)
    paragraph_border_bottom(rule, ACCENT, "24", "8")

    metadata = [
        ("Organization", "Kreyco Tech Support"),
        ("System", "Public Google Apps Script portal integrated with monday.com"),
        ("Version", "1.1"),
        ("Current-state date", "September 3, 2026"),
        ("Audience", "Tech Support, system administrators, coaches, developers, and AI assistants"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.75)
        p.paragraph_format.right_indent = Inches(0.75)
        r = p.add_run(label + ": ")
        set_font(r, size=10.5, color=ACCENT_DARK, bold=True)
        r = p.add_run(value)
        set_font(r, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, ACCENT_LIGHT, ACCENT)
    r = p.add_run("Internal operational documentation. Do not add API tokens, passwords, portal signing secrets, reusable credentials, or signed coach-link tokens to this document or to an AI prompt.")
    set_font(r, size=9.5, color=ACCENT_DARK, bold=True)

    doc.add_page_break()


def add_contents(doc, source_lines):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    seen = []
    for line in source_lines:
        if line.startswith("## "):
            title = line[3:].strip()
            if title.lower().startswith("system knowledge"):
                continue
            seen.append(title)
    for title in seen:
        p = doc.add_paragraph(style="Contents Entry")
        add_inline(p, title, size=10.5)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    add_inline(note, "This is a static contents list for reliable Google Docs and PDF ingestion. Word heading navigation remains available in the editable file.", size=9.5, color=MUTED)
    doc.add_page_break()


def render_markdown(doc, lines, bullet_abs, decimal_abs):
    first_rule_seen = False
    in_code = False
    code_lines = []
    bullet_num = None
    decimal_num = None
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if stripped.startswith("# ") or stripped.startswith("## System knowledge") or stripped.startswith("**Organization:") or stripped.startswith("**System:") or stripped.startswith("**Document version:") or stripped.startswith("**Current-state date:") or stripped.startswith("**Audience:") or stripped.startswith("**Classification:"):
            idx += 1
            continue
        if stripped == "---":
            first_rule_seen = True
            idx += 1
            continue
        if not first_rule_seen:
            idx += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                p.paragraph_format.keep_together = True
                shade_paragraph(p, SOFT, LINE)
                r = p.add_run("\n".join(code_lines))
                set_font(r, name="Courier New", size=8.5, color=INK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            bullet_num = None
            decimal_num = None
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            bullet_num = None
            decimal_num = None
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:].strip(), size=16, color=ACCENT_DARK)
            for run in p.runs:
                run.bold = True
            bullet_num = None
            decimal_num = None
            idx += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip(), size=13, color=ACCENT_DARK)
            for run in p.runs:
                run.bold = True
            bullet_num = None
            decimal_num = None
            idx += 1
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[5:].strip(), size=12, color=ACCENT_DARK)
            for run in p.runs:
                run.bold = True
            bullet_num = None
            decimal_num = None
            idx += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if ordered:
            if decimal_num is None:
                decimal_num = new_num_id(doc, decimal_abs)
            p = doc.add_paragraph()
            apply_numbering(p, decimal_num)
            add_inline(p, ordered.group(2))
            bullet_num = None
            idx += 1
            continue
        if bullet:
            if bullet_num is None:
                bullet_num = new_num_id(doc, bullet_abs)
            p = doc.add_paragraph()
            apply_numbering(p, bullet_num)
            add_inline(p, bullet.group(1))
            decimal_num = None
            idx += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        bullet_num = None
        decimal_num = None
        idx += 1


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("CLASSROOM CREATION REQUEST  |  SYSTEM HANDBOOK")
    set_font(r, size=8, color=MUTED, bold=True)
    paragraph_border_bottom(p, LINE, "6", "3")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Kreyco Tech Support  |  Internal  |  Page ")
    set_font(r, size=8, color=MUTED)
    r = p.add_run()
    set_font(r, size=8, color=MUTED)
    add_field(r, "PAGE")

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kreyco Tech Support  |  Internal operational documentation")
    set_font(r, size=8, color=MUTED)


def set_document_options(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.core_properties.title = "Classroom Creation Request - System Knowledge Base and Operations Handbook"
    doc.core_properties.subject = "Google Apps Script and monday.com classroom request portal"
    doc.core_properties.author = "Kreyco Tech Support"
    doc.core_properties.keywords = "Classroom Creation Request, monday.com, Google Apps Script, Gemini, operations"
    doc.core_properties.comments = "Generated from the repository's canonical knowledge-base Markdown."


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_knowledge_base.py INPUT.md OUTPUT.docx")
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    output.parent.mkdir(parents=True, exist_ok=True)
    lines = source.read_text(encoding="utf-8").splitlines()

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    set_document_options(doc)
    bullet_abs, decimal_abs = configure_numbering(doc)
    build_cover(doc)
    add_contents(doc, lines)
    render_markdown(doc, lines, bullet_abs, decimal_abs)

    # Avoid widows/orphans and keep table rows intact where possible.
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    doc.save(output)


if __name__ == "__main__":
    main()
